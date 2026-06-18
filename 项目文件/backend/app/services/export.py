"""导出服务：Word / PDF / Excel"""

from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font

from app.models.record import Record


def _format_field_value(field_type: str, value: Any) -> str:
    """根据字段类型格式化值为可读字符串。"""
    if value is None:
        return ""

    if field_type in ("text", "textarea"):
        return str(value)

    if field_type == "richtext":
        # 提取纯文本：去除 HTML 标签
        text = re.sub(r"<[^>]+>", "", str(value))
        return text.strip()

    if field_type == "number":
        return str(value)

    if field_type == "datetime":
        if isinstance(value, str):
            return value
        if isinstance(value, datetime):
            return value.strftime("%Y-%m-%d %H:%M")
        return str(value)

    if field_type == "select":
        if isinstance(value, list):
            return ", ".join(str(v) for v in value)
        return str(value)

    if field_type == "multiselect":
        if isinstance(value, list):
            return ", ".join(str(v) for v in value)
        return str(value)

    if field_type == "boolean":
        return "是" if value else "否"

    if field_type in ("file", "image"):
        if isinstance(value, dict):
            return value.get("name", value.get("filename", str(value)))
        if isinstance(value, list):
            names = []
            for item in value:
                if isinstance(item, dict):
                    names.append(item.get("name", item.get("filename", str(item))))
                else:
                    names.append(str(item))
            return ", ".join(names)
        return str(value)

    # divider 及其他未知类型
    return str(value)


def _extract_fields(record: Record) -> list[tuple[str, str]]:
    """从记录中提取 (label, formatted_value) 列表，跳过 divider。"""
    snapshot = record.template_snapshot if isinstance(record.template_snapshot, list) else []
    data: dict = record.data or {}
    fields = []
    for field_def in snapshot:
        field_type = field_def.get("type", "text")
        if field_type == "divider":
            continue
        key = field_def.get("key", "")
        label = field_def.get("label", key)
        value = data.get(key)
        formatted = _format_field_value(field_type, value)
        fields.append((label, formatted))
    return fields


def export_to_word(record: Record) -> BytesIO:
    """使用 python-docx 生成 Word 文档。"""
    doc = Document()

    # 标题
    title_para = doc.add_heading(record.title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 模板信息
    template_name = "未知模板"
    template_version = ""
    if record.template:
        template_name = record.template.name
        template_version = str(record.template.version)
    elif record.template_snapshot:
        # 从 snapshot 中无法直接获取模板名称，留空
        pass

    info_para = doc.add_paragraph()
    info_para.add_run("模板：").bold = True
    info_para.add_run(template_name)
    if template_version:
        info_para.add_run(f"  |  版本：{template_version}")

    doc.add_paragraph("")  # 空行

    # 字段表格
    fields = _extract_fields(record)
    if fields:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = "字段"
        header_cells[1].text = "值"
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        # 数据行
        for label, value in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_to_pdf(record: Record) -> BytesIO:
    """使用 WeasyPrint 生成 PDF。"""
    try:
        from weasyprint import HTML  # type: ignore[import-untyped]
    except (ImportError, OSError) as e:
        raise RuntimeError(
            "WeasyPrint 依赖的系统库未安装（pango/gobject）。"
            "请参考 https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation"
        ) from e

    template_name = "未知模板"
    template_version = ""
    if record.template:
        template_name = record.template.name
        template_version = str(record.template.version)

    fields = _extract_fields(record)

    # 构建字段行 HTML
    rows_html = ""
    for label, value in fields:
        ev = value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        el = label.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        rows_html += f"""
        <tr>
            <td class="label">{el}</td>
            <td class="value">{ev}</td>
        </tr>"""

    version_info = f"  |  版本：{template_version}" if template_version else ""
    meta_line = f"模板：{template_name}{version_info}"

    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        @page {{ size: A4; margin: 2cm; }}
        body {{ font-family: sans-serif; font-size: 12pt; color: #333; }}
        h1 {{ text-align: center; font-size: 18pt; margin-bottom: 10px; }}
        .meta {{ text-align: center; color: #666; font-size: 10pt; margin-bottom: 20px; }}
        table {{ width: 100%; border-collapse: collapse; }}
        th, td {{ border: 1px solid #ccc; padding: 8px 12px; text-align: left; }}
        th {{ background-color: #f5f5f5; font-weight: bold; }}
        td.label {{ width: 30%; font-weight: bold; background-color: #fafafa; }}
        td.value {{ width: 70%; }}
    </style>
</head>
<body>
    <h1>{record.title}</h1>
    <div class="meta">{meta_line}</div>
    <table>
        <thead>
            <tr><th>字段</th><th>值</th></tr>
        </thead>
        <tbody>
            {rows_html}
        </tbody>
    </table>
</body>
</html>"""

    buf = BytesIO()
    HTML(string=html_content).write_pdf(buf)
    buf.seek(0)
    return buf


def export_to_excel(record: Record) -> BytesIO:
    """使用 openpyxl 生成 Excel。"""
    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = record.title[:31]  # Excel 工作表名最长 31 字符

    # 表头加粗
    header_font = Font(bold=True)
    ws["A1"] = "字段"
    ws["B1"] = "值"
    ws["A1"].font = header_font
    ws["B1"].font = header_font

    # 数据行
    fields = _extract_fields(record)
    for idx, (label, value) in enumerate(fields, start=2):
        ws[f"A{idx}"] = label
        ws[f"B{idx}"] = value

    # 调整列宽
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 60

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf
